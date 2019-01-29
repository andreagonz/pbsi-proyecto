from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.decorators import login_required
from django.views.generic import View
from django.urls import reverse_lazy
from django.shortcuts import render
from phishing.aux import log
from phishing.forms import GraficasForm
from phishing.views import aux
from phishing.models import Url, ASN
import randomcolor
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
import numpy as np
import matplotlib.pyplot as plt
from django.db.models import Count, Q, F, Avg
import datetime
from django.http import HttpResponse
import random
import string
import os

class DocumentView(LoginRequiredMixin, View):
    
    def get(self,request, *args, **kwargs):
            return render(request,'reporte/reporte.html', {'form': GraficasForm()})

def agrega_imagen(fig, documento, titulo, descripcion):
    try:
        a = ''.join(random.choice(string.ascii_uppercase + string.digits) for _ in range(8))
        path = '/tmp/%s.png' % a
        fig.savefig(path)
        plt.clf()
        q = documento.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.add_run(titulo.upper()).bold = True
        q.add_run("\n%s" % descripcion)
        documento.add_picture(path)
        q = documento.add_paragraph()
        run = q.add_run()
        run.add_break(WD_BREAK.PAGE)
        os.remove(path)
    except Exception as e:
        log.log('Error: %s' % str(e), 'reportes.log')

def url_info(u, q, d):
    i = u.obten_info
    if i and i.captura and hasattr(i.captura, 'file'):
        d.add_picture(i.captura.file, width=Inches(4.0))
        q = d.add_paragraph("")
    q.add_run("URL: %s\n" % u.url)
    q.add_run("IP: %s\n" % u.dominio.ip_str)
    q.add_run("Código: %s\n" % u.codigo_str)
    q.add_run("Fecha de creación: %s\n" % u.timestamp_creacion)    
    q.add_run("Detección: %s\n" % u.deteccion_str)
    q.add_run("Estado: %s\n" % u.estado)
    q.add_run("Entidad: %s\n" % u.entidad_afectada_str)
    q.add_run("Correos: %s\n" % u.dominio.correos_str)
    q.add_run("ISP: %s\n" % u.dominio.isp_str)
    q.add_run("País: %s\n" % u.dominio.pais_str)
    q.add_run("ASN: %s\n" % u.dominio.asn_str)
    q.add_run("Servidor: %s\n" % u.dominio.servidor_str)
    q.add_run("RIR: %s\n" % u.dominio.rir_str)
    q.add_run("Servidores DNS: %s\n" % u.dominio.dns_str)    
    if i:
        q.add_run("Título: %s\n" % i.titulo_str)
        q.add_run("Ofuscacion: %s\n" % i.ofuscaciones_str)
        if i.hash_archivo:
            q.add_run("Hash MD5 de archivo: %s\n" % i.hash_archivo)
    if u.es_redireccion:
        red = u.obten_info_redireccion
        if red:
            if red.redireccion:
                q.add_run("Redirección: %s\n" % red.redireccion.url)
            if red.redireccion_final:
                q.add_run("Redirección final: %s\n" % red.redireccion_final.url)
                     
@login_required(login_url=reverse_lazy('login'))
def crear_doc(request):
    if request.method == 'POST':
        form = GraficasForm(request.POST)
        if form.is_valid():
            archivo = form.cleaned_data['archivo']
            sitios = form.cleaned_data['sitios']
            top_sitios = form.cleaned_data['top_sitios']
            sectores = form.cleaned_data['sectores']
            entidades = form.cleaned_data['entidades']
            detecciones = form.cleaned_data['detecciones']
            tiempo_reporte = form.cleaned_data['tiempo_reporte']
            top_paises = form.cleaned_data['top_paises']
            top_hosting = form.cleaned_data['top_hosting']

            sitios_info = form.cleaned_data['sitios_info']
            top_sitios_info = form.cleaned_data['top_sitios_info']
            sectores_info = form.cleaned_data['sectores_info']
            entidades_info = form.cleaned_data['entidades_info']
            detecciones_info = form.cleaned_data['detecciones_info']
            tiempo_reporte_info = form.cleaned_data['tiempo_reporte_info']
            top_paises_info = form.cleaned_data['top_paises_info']
            top_hosting_info = form.cleaned_data['top_hosting_info']
            
            urls_info = form.cleaned_data['urls']

            graficas = []
            rand_color = randomcolor.RandomColor()
            document = Document()
            document.add_heading('Reporte', 0)
            p = document.add_paragraph('Reporte elaborado por la herramienta ')
            p.add_run('SAAPM').bold = True            
            inicio = form.cleaned_data['inicio']
            fin = form.cleaned_data['fin']
            
            urls0 = Url.objects.filter(timestamp_creacion__lte=fin, timestamp_creacion__gte=inicio)
            lst = []
            for u in urls0:
                d = u.obten_info
                if d and (d.deteccion == 'P' or d.deteccion == 'I'):
                    lst.append(u.pk)
            urls = Url.objects.filter(pk__in=lst)
                
            document.add_heading('Periodo',level=1)
            q = document.add_paragraph('De: ')
            q.add_run(str(inicio)).bold = True
            q.add_run('      ')
            q.add_run('A :  ')
            q.add_run(str(fin)).bold = True

            if sitios:
                sitios_activos = urls.filter(timestamp_desactivado__lt=fin)
                sitios_reportados = urls.filter(ticket__timestamp__lt=fin)
                x = ['Activos', 'Reportados', 'Detectados']
                y = [sitios_activos.count(), sitios_reportados.count(), urls.count()]
                y_pos = np.arange(len(x))
                fig, ax = plt.subplots()
                ax.set_ylabel('Número de sitios')
                ax.bar(y_pos, y, align='center', alpha=0.5)
                plt.xticks(y_pos, x)
                titulo = 'Estados de sitios phishing'
                agrega_imagen(fig, document, titulo, sitios_info)

            if top_sitios:
                top_sitios = urls.filter(
                    Q(timestamp_desactivado__isnull=True) | Q(timestamp_desactivado__lt=fin)
                ).annotate(tiempo_vida=(fin - F('timestamp_creacion'))).order_by('-tiempo_vida')[:5]
                y = [x.url for x in top_sitios]
                x = [aux.delta_horas(x.tiempo_vida) for x in top_sitios]                
                y_pos = np.arange(len(x))
                fig, ax = plt.subplots()
                fig.subplots_adjust(left=0.5)
                ax.set_xlabel('T (Horas)')
                ax.barh(y_pos, x, align='center', alpha=0.5)
                plt.yticks(y_pos, y)
                titulo = 'Top 5 – Sitios phishing vs Tiempo de vida'
                agrega_imagen(fig, document, titulo, top_sitios_info)
                
            if sectores:
                sectores = urls.values('urlactiva__entidad_afectada__clasificacion__nombre').annotate(
                    cuenta=Count('urlactiva__entidad_afectada__clasificacion__nombre'))
                x = [x['urlactiva__entidad_afectada__clasificacion__nombre']
                     for x in sectores
                     if x['urlactiva__entidad_afectada__clasificacion__nombre']]
                y = [x['cuenta'] for x in sectores
                     if x['urlactiva__entidad_afectada__clasificacion__nombre']]
                colores = rand_color.generate(count=len(x))
                fig, ax = plt.subplots()
                ax.pie(y, labels=x, colors=colores, autopct='%1.1f%%', startangle=90)
                titulo = 'Sectores afectados'
                ax.axis('equal')
                agrega_imagen(fig, document, titulo, sectores_info)
                
            if entidades:
                entidades = urls.values('urlactiva__entidad_afectada__nombre').annotate(
                    cuenta=Count('urlactiva__entidad_afectada__nombre'))
                x = [x['urlactiva__entidad_afectada__nombre'] for x in entidades
                        if x['urlactiva__entidad_afectada__nombre']]
                y = [x['cuenta'] for x in entidades
                        if x['urlactiva__entidad_afectada__nombre']]
                colores = rand_color.generate(count=len(x))
                fig, ax = plt.subplots()
                ax.pie(y, labels=x, colors=colores, autopct='%1.1f%%', startangle=90)
                titulo = 'Entidades afectadas'
                ax.axis('equal')
                agrega_imagen(fig, document, titulo, entidades_info)

            if detecciones:
                ndias = (fin - inicio).days
                fechas = [inicio + datetime.timedelta(days=i) for i in range(ndias + 1)]
                y = []
                for d in fechas:
                    y.append(urls.filter(timestamp_creacion__date=d.date()).count())
                x = [str(f) for f in fechas]
                y_pos = np.arange(len(x))
                fig, ax = plt.subplots()
                fig.subplots_adjust(bottom=0.2)
                ax.set_ylabel('Número de detecciones')
                ax.bar(y_pos, y, align='center', alpha=0.5)
                plt.xticks(y_pos, x, rotation=45)
                titulo = 'Número de detecciones por fecha'
                agrega_imagen(fig, document, titulo, detecciones_info)

            if tiempo_reporte:
                ndias = (fin - inicio).days
                fechas = [inicio + datetime.timedelta(days=i) for i in range(ndias + 1)]
                y = []
                x = [str(f) for f in fechas]
                tiempo_promedio_reporte = []
                tiempo_promedio_postreporte = []
                for d in fechas:
                    sitios = urls.filter(
                        timestamp_creacion__date=d
                    ).exclude(ticket__isnull=True)
                    tiempo_promedio_reporte.append(sitios.annotate(
                        tiempo_reportado=F('ticket__timestamp') - F('timestamp_creacion')).aggregate(
                            Avg('tiempo_reportado')).get('tiempo_reportado__avg', 0))
                    tiempo_promedio_postreporte.append(sitios.filter(
                        timestamp_desactivado__lte=fin).annotate(
                            tiempo_reportado=F('timestamp_desactivado') - F('ticket__timestamp')
                        ).aggregate(Avg('tiempo_reportado')).get('tiempo_reportado__avg', 0))
                y1 = [aux.delta_horas(x) if x else 0 for x in tiempo_promedio_reporte]
                y2 = [aux.delta_horas(x) if x else 0 for x in tiempo_promedio_postreporte]
                fig, ax = plt.subplots()
                fig.subplots_adjust(bottom=0.3)
                line1, = ax.plot(x, y1, linewidth=2,
                                 label='Tiempo promedio de reporte')
                line2, = ax.plot(x, y2, linewidth=2,
                                 label='Tiempo promedio de vida postreporte')
                plt.xticks(rotation=45)
                ax.set_ylabel('T (Horas)')
                titulo = 'Tiempo promedio de reporte vs tiempo promedio de vida postreporte'
                ax.legend(loc='best')
                agrega_imagen(fig, document, titulo, tiempo_reporte_info)

            if top_paises:                
                paises = urls.exclude(dominio__pais__isnull=True).values(
                    'dominio__pais').annotate(
                        cuenta=Count('dominio__pais')).order_by('-cuenta')[:10]
                x = [a['dominio__pais'] for a in paises]
                y = [a['cuenta'] for a in paises]
                fig, ax = plt.subplots()
                ax.set_ylabel('Número de sitios')
                y_pos = np.arange(len(x))
                ax.bar(y_pos, y, align='center', alpha=0.5)
                plt.xticks(y_pos, x)
                titulo = 'Top 10 países que hospedan phishing'
                agrega_imagen(fig, document, titulo, top_paises_info)
                
            if top_hosting:
                hosting = urls.exclude(dominio__asn__isnull=True).values(
                    'dominio__asn').annotate(
                        cuenta=Count('dominio__asn')).order_by('-cuenta')[:10]
                asns = ASN.objects.filter(pk__in=[x['dominio__asn'] for x in hosting])
                l = []
                for a in asns:
                    l.append((str(a), next(x['cuenta'] for x in hosting if
                                           x['dominio__asn'] == a.pk)))
                l.sort(key=lambda x:x[1], reverse=True)
                x = [a[0] for a in l]
                y = [a[1] for a in l]                
                fig, ax = plt.subplots()
                fig.subplots_adjust(bottom=0.5)
                ax.set_ylabel('Número de sitios')
                y_pos = np.arange(len(x))
                ax.bar(y_pos, y, align='center', alpha=0.5)
                plt.xticks(y_pos, x, rotation=70)
                titulo = 'Top 10 servicios de hosting que hospedan phishing'
                agrega_imagen(fig, document, titulo, top_hosting_info)
                
            if urls_info:
                q = document.add_paragraph('')
                q.add_run("INFORMACIÓN SOBRE URLS:\n").bold = True
                for u in urls0:
                    url_info(u, q, document)

            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = 'attachment; filename=%s.docx' % archivo
            document.save(response)
            return response
        else:
            return render(request,'reporte/reporte.html', {'form': GraficasForm()})
