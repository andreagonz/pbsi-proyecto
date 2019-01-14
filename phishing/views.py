from django.shortcuts import render, redirect, get_object_or_404
from django.forms import Textarea
from django.urls import reverse_lazy
from django.contrib.auth.decorators import login_required
from .forms import *
from .models import *
from .phishing import *
from .correo import *
from django.views.generic import TemplateView
from django.template import loader
from django.views.generic.detail import DetailView
from django.http import HttpResponse, Http404
from django.contrib.postgres.search import SearchQuery, SearchRank, SearchVector
from django.core.exceptions import MultipleObjectsReturned
from django.conf import settings
from shutil import copyfile
import os
from .reporte import *
from datetime import timedelta, datetime
from django.utils import timezone
from time import mktime
import time
from django.urls import reverse_lazy
from django.views.generic.edit import UpdateView, CreateView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.db.models import Avg
from django.views.generic import TemplateView,View
from django.template import RequestContext
from django.http import HttpResponse
from phishing.forms import *
from django.core.exceptions import MultipleObjectsReturned
from django.http import JsonResponse
from django.db.models import Count, Q,F
from django.db.models.functions import Extract
import json
from rest_framework.views import APIView
from rest_framework.response import Response
import randomcolor
import datetime
from phishing.phishing import lineas_md5, md5, archivo_hashes, leeComentariosHTML, archivo_texto
from docx import Document
from docx.shared import Inches
from .entrada import( lee_csv, lee_txt, lee_json )
import matplotlib.pyplot as plt
import tempfile
import random
import string
import numpy as np
from docx.shared import Inches

def log_x(mensaje, bitacora):
    t = timezone.localtime(timezone.now())
    l = os.path.join(settings.DIR_LOG, bitacora)
    with open(l, 'a') as w:
        w.write('[%s] %s\n' % (t, mensaje))

@login_required(login_url=reverse_lazy('login'))
def monitoreo(request):
    dominios = Dominio.objects.all()
    activos = [x for x in dominios if x.activo]
    return render(request, 'monitoreo.html', context={'dominios':activos})

def rmimg(img):
    if img is None:
        return
    f = os.path.join(settings.MEDIA_ROOT, img)
    if os.path.exists(f):
        os.remove(f)

def cpimg(img, img2):
    if img is None or img2 is None:
        return
    f = os.path.join(settings.BASE_DIR, img[1:])
    f2 = os.path.join(settings.BASE_DIR, img2[1:])
    if os.path.exists(f):
        copyfile(f, f2)
    
def url_reporta(url, ts):
    url.reportado = True
    if url.deteccion == 'I':
        url.timestamp_deteccion = ts
        url.deteccion = 'P'
    url.save()

def url_ignora(url):
    url.ignorado = True
    url.deteccion = 'N'
    url.save()

@login_required(login_url=reverse_lazy('login'))
def monitoreo_id(request, pk):
    dominio = get_object_or_404(Dominio, pk=pk)
    if not dominio.activo:
        raise Http404()
    urls = dominio.urls_activas
    context = {
        'dominio': dominio,
        'monitoreo': False,
        'activo': True,
    }
    proxy_form = ProxyForm()
    hoy = timezone.localtime(timezone.now())
    md = md5(dominio.dominio.encode('utf-8', 'backslashreplace'))
    ticket = ('%d%02d%02d%s' % (hoy.year, hoy.month, hoy.day, md[:7])).upper()
    correos = []
    for url in urls:
        for x in url.dominio.correos.all():
            correos.append(str(x))
    correos = list(set(correos))
    datos = {
        'de': settings.CORREO_DE,
        'para': ', '.join(correos),
        'asunto': obten_asunto(dominio, ticket),
        'mensaje': obten_mensaje(dominio, ticket)
    }
    mensaje_form = MensajeForm(initial=datos, urls=urls)
    if request.method == 'POST':
        if request.POST.get('boton-curl'):
            proxy_form = ProxyForm(request.POST)
            if proxy_form.is_valid():
                http = proxy_form.cleaned_data['http']
                https = proxy_form.cleaned_data['https']
                tor = proxy_form.cleaned_data['tor']
                proxies = proxy_form.cleaned_data['proxy']
                proxy = None
                if tor:
                    proxy = {'http':  'socks5://127.0.0.1:9050', 'https': 'socks5://127.0.0.1:9050'}
                elif http or https:
                    proxy = {}
                    if http:
                        proxy['http'] = http
                    if https:
                        proxy['https'] = https
                elif not proxies is None and (not proxies.http is None or
                                              not proxies.http is None):
                    proxy = {}
                    if not proxies.http is None:
                        proxy['http'] = proxies.http
                    if not proxies.https is None:
                        proxy['https'] = proxies.https
                monitorea_dominio(dominio.dominio, dominio.urls_activas, proxy)
            context['monitoreo'] = True
            mensaje_form.actualiza()
            context['mensaje_form'] = mensaje_form
            context['activo'] = dominio.activo
            context['proxy_form'] = proxy_form
            return render(request, 'monitoreo_id.html', context)
        elif request.POST.get('boton-mensaje'):
            # if not dominio.activo:
            # return render(request, 'monitoreo_error.html', {'dominio': dominio})
            mensaje_form = MensajeForm(request.POST, urls=urls)
            if mensaje_form.is_valid():
                if not mensaje_form.cleaned_data.get('para', None):
                    if not mensaje_form._errors.get('para', None):
                        from django.forms.utils import ErrorList
                        mensaje_form._errors['para'] = ErrorList()
                    mensaje_form._errors['para'].append('Campo necesario')
                    context['mensaje_form'] = mensaje_form
                    context['proxy_form'] = proxy_form
                    return render(request, 'monitoreo_id.html', context)
                de = mensaje_form.cleaned_data['de']
                para = [x.strip() for x in mensaje_form.cleaned_data['para'].split(',')]
                cc = [x.strip() for x in mensaje_form.cleaned_data['cc'].split(',')]
                cco = [x.strip() for x in mensaje_form.cleaned_data['cco'].split(',')]
                asunto = mensaje_form.cleaned_data['asunto']
                mensaje = mensaje_form.cleaned_data['mensaje']
                capturas = mensaje_form.cleaned_data['capturas']
                urls_reportadas = mensaje_form.cleaned_data['urls']
                msg = genera_mensaje(dominio, de, para, cc, cco, asunto, mensaje, capturas)
                enviado = manda_correo(para, cc, cco, msg)
                ts = timezone.localtime(timezone.now())
                if not enviado:
                    context['dominio'] = dominio
                    return render(request, 'mensaje_error.html', context)
                try:
                    men = Mensaje.objects.get(ticket=ticket)
                except:
                    men = Mensaje(ticket=ticket)
                    men.save()
                men.timestamp = ts
                tsd = url.timestamp_deteccion if url.timestamp_deteccion else ts
                for x in urls_reportadas:
                    mu = MensajeURL(mensaje=men, timestamp_creacion_sitio=x.timestamp_reactivacion,
                                    url=x, timestamp_deteccion=tsd)
                    mu.pais = x.dominio.pais
                    mu.asn = x.dominio.asn
                    mu.save()
                    mu.entidades_afectadas.add(*x.entidades_afectadas.all())
                    url_reporta(x, ts)
                men.save()
                context = {
                    'dominio': dominio,
                    'urls': ', '.join([url.url for x in urls_reportadas]),
                    'de': de,
                    'para': ', '.join(para),
                    'cc': ', '.join(cc),
                    'cco': ', '.join(cco),
                    'asunto': asunto,
                    'mensaje': mensaje,
                    'capturas': capturas
                }
                return render(request, 'monitoreo_exito.html', context)
        elif request.POST.get('boton-ignorar') and request.user.is_superuser:
            for x in urls:
                url_ignora(x)
            return redirect('monitoreo')
    context['mensaje_form'] = mensaje_form
    context['proxy_form'] = proxy_form
    return render(request, 'monitoreo_id.html', context)

def context_reporte(sitios):
    activas = urls_activas(sitios)
    inactivas = urls_inactivas(sitios)
    redirecciones = urls_redirecciones(sitios)
    context = {
        'urls_total': cuenta_urls(sitios),
        'num_urls_activas': len(set([x.url for x in activas])),
        'num_urls_inactivas': len(set([x.url for x in inactivas])),
        'num_urls_redirecciones': len(set([x.url for x in redirecciones])),
        'entidades': urls_entidades(sitios),
        'titulos': urls_titulos(sitios),
        'dominios': urls_dominios(sitios),
        'paises': urls_paises(sitios),
        'activas': activas,
        'inactivas': inactivas,
        'redirecciones': redirecciones
    }
    return context

def valida_urls(request):
    if request.method == 'POST':
        if request.POST.get("boton_urls"):
            form = UrlsForm(request.POST) 
            if form.is_valid():                
                urls = form.cleaned_data['urls']
                urls_limpias = []
                for x in urls.split('\n'):
                    x = x.strip()
                    if x:
                        for y in x.split(','):
                            y = y.strip()
                            if y:
                                urls_limpias.append(y)
                sitios = verifica_urls(urls_limpias, None, False)
                no_reportados = False
                for x in sitios:
                    if not x.reportado:
                        no_reportados = True
                context = context_reporte(sitios)
                context['no_reportados'] = no_reportados
                return render(request, 'reporte_urls.html', context)
        elif request.POST.get("boton_archivo") and request.FILES.get('file', None):
            form = ArchivoForm(request.POST)
            f = request.FILES['file'].read().decode('utf-8')
            name = request.FILES['file'].name
            urls = []
            if name.endswith('.txt'):
                urls = lee_txt(f)
            elif name.endswith('.json'):
                urls = lee_json(f)
            elif name.endswith('.csv'):
                urls = lee_csv(f)
            sitios = verifica_urls(urls, None, False)
            no_reportados = False
            for x in sitios:
                if not x.reportado:
                    no_reportados = True
            context = context_reporte(sitios)
            context['no_reportados'] = no_reportados
            return render(request, 'reporte_urls.html', context)
    form1 = UrlsForm()
    form2 = ArchivoForm()
    return render(request, 'valida_urls.html', {'form1': form1, 'form2': form2})

def url_detalle(request, pk):
    url = get_object_or_404(Url, pk=pk)
    # comentarios = archivo_comentarios(url)
    hashes = archivo_hashes(url)
    context = {
        'url': url,
        # 'comentarios': comentarios
        'hashes': hashes
    }
    return render(request, 'url_detalle.html', context)

@login_required(login_url=reverse_lazy('login'))
def busca(request):
    message = "No se encontraron coincidencias"
    resultados = []
    q = ''
    if request.method == "GET" and request.GET.get('q', None):
        q = request.GET['q'].strip()
        """
        if request.GET.get('a', None) and request.GET['a'].strip() == '1':
            urls = Url.objects.exclude(archivo=None)
            for u in urls:
                f = archivo_texto(u)
                for h in lineas_md5(f):
                    if h == q:
                        resultados.append(u)
                for x in leeComentariosHTML(f):
                    if q in x:
                        resultados.append(u)
                if u.hash_archivo == q:
                    resultados.append(u)
            resultados = list(set(resultados))
        """
        resultados = Url.objects.annotate(
            search=SearchVector('url') + SearchVector('dominio__dominio') +
            SearchVector('entidades_afectadas__clasificacion__nombre') +
            SearchVector('entidades_afectadas__nombre') +
            SearchVector('ofuscacion__nombre') + SearchVector('dominio__correos__correo') +
            SearchVector('dominio__rir__nombre') + SearchVector('dominio__dns__nombre') +
            SearchVector('dominio__ip') + SearchVector('dominio__pais') +
            SearchVector('dominio__asn') + SearchVector('dominio__isp') +
            SearchVector('identificador') + SearchVector('titulo') +
            SearchVector('hash_archivo') + SearchVector('redireccion') +
            SearchVector('dominio__servidor') +
            SearchVector('mensaje__ticket')
        ).filter(search=SearchQuery(q)).distinct('url')            
    return render(request, 'results.html',
                  {'resultados': resultados,
                   'query': q
                  })

@login_required(login_url=reverse_lazy('login'))
def muestraResultados(request,srch):
	return render(request,'results.html',{})
    
@login_required(login_url=reverse_lazy('login'))
def historico(request):
    fin = timezone.localtime(timezone.now()).date()
    inicio = fin
    form = HistoricoForm()
    if request.method == 'POST':
        form = HistoricoForm(request.POST)
        if form.is_valid():
            inicio = form.cleaned_data['inicio']
            fin = form.cleaned_data['fin']
    sitios = Url.objects.filter(timestamp_reactivacion__date__lte=fin,
                                timestamp_reactivacion__date__gte=inicio)
    context = context_reporte(sitios)
    context['inicio'] = inicio
    context['fin'] = fin
    context['form'] = form
    return render(request, 'historico.html', context)

@login_required(login_url=reverse_lazy('login'))
def ajustes(request):
    proxies = Proxy.objects.all()
    recursos = Recurso.objects.all()
    asunto_form = CambiaAsuntoForm(initial={'asunto': lee_plantilla_asunto()})
    mensaje_form = CambiaMensajeForm(initial={'mensaje': lee_plantilla_mensaje()})
    unam_asunto_form = CambiaUnamAsuntoForm(initial={'asunto': lee_plantilla_unam_asunto()})
    unam_mensaje_form = CambiaUnamMensajeForm(initial={'mensaje': lee_plantilla_unam_mensaje()})
    actualizacion_form = FrecuenciaForm()
    verificacion_form = FrecuenciaForm()
    if request.method == 'POST':
        if request.POST.get('cambia-asunto'):
            asunto_form = CambiaAsuntoForm(request.POST)
            if asunto_form.is_valid():
                asunto = asunto_form.cleaned_data['asunto']
                cambia_asunto(asunto)
        elif request.POST.get('cambia-mensaje'):
            mensaje_form = CambiaMensajeForm(request.POST)
            if mensaje_form.is_valid():
                mensaje = mensaje_form.cleaned_data['mensaje']
                cambia_mensaje(mensaje)
        elif request.POST.get('cambia-unam-asunto'):
            unam_asunto_form = CambiaUnamAsuntoForm(request.POST)
            if unam_asunto_form.is_valid():
                asunto = unam_asunto_form.cleaned_data['asunto']
                cambia_unam_asunto(asunto)
        elif request.POST.get('cambia-unam-mensaje'):
            unam_mensaje_form = CambiaUnamMensajeForm(request.POST)
            if unam_mensaje_form.is_valid():
                mensaje = unam_mensaje_form.cleaned_data['mensaje']
                cambia_unam_mensaje(mensaje)
        elif request.POST.get('cambia-actualizacion'):
            actualizacion_form = FrecuenciaForm(request.POST)
            if actualizacion_form.is_valid():
                actualizacion = actualizacion_form.cleaned_data['frecuencia']
                if actualizacion < 1:
                    actualizacion = 8
                cambia_frecuencia('actualiza', actualizacion)
        elif request.POST.get('cambia-verificacion'):
            verificacion_form = FrecuenciaForm(request.POST)
            if verificacion_form.is_valid():
                verificacion = verificacion_form.cleaned_data['frecuencia']
                if verificacion < 1:
                    verificacion = 1
                cambia_frecuencia('verifica', verificacion)
    context = {
        'recursos': recursos,
        'proxies': proxies,
        'asunto_form': asunto_form,
        'mensaje_form': mensaje_form,
        'asunto_unam_form': unam_asunto_form,
        'mensaje_unam_form': unam_mensaje_form,
        'actualizacion_form': actualizacion_form,
        'verificacion_form': verificacion_form,
    }
    return render(request, 'ajustes.html', context)

@login_required(login_url=reverse_lazy('login'))
def elimina_proxy(request, pk):
    proxy = get_object_or_404(Proxy, pk=pk)
    proxy.delete()
    return redirect('ajustes')

class ActualizaProxy(LoginRequiredMixin, UpdateView):
    model = Proxy
    template_name = 'actualiza_proxy.html'
    success_url = reverse_lazy('ajustes')
    fields = ('http', 'https')
    
class NuevoProxy(LoginRequiredMixin, CreateView):
    model = Proxy
    template_name = 'nuevo_proxy.html'
    success_url = reverse_lazy('ajustes')
    fields = ('http', 'https')

@login_required(login_url=reverse_lazy('login'))
def elimina_recurso(request, pk):
    recurso = get_object_or_404(Recurso, pk=pk)
    recurso.delete()
    return redirect('ajustes')

class ActualizaRecurso(LoginRequiredMixin, UpdateView):
    model = Recurso
    template_name = 'actualiza_recurso.html'
    success_url = reverse_lazy('ajustes')
    fields = ('es_phishtank', 'recurso', 'max_urls')
    
class NuevoRecurso(LoginRequiredMixin, CreateView):
    model = Recurso
    template_name = 'nuevo_recurso.html'
    success_url = reverse_lazy('ajustes')
    fields = ('es_phishtank', 'recurso', 'max_urls')

@login_required(login_url=reverse_lazy('login'))
def ofuscaciones_view(request):
    of = Ofuscacion.objects.all()
    context = {
        'ofuscaciones': of
    }
    return render(request, 'ofuscaciones.html', context)

@login_required(login_url=reverse_lazy('login'))
def entidades_view(request):
    context = {
        'clasificaciones': Clasificacion_entidad.objects.all(),
        'entidades': Entidades.objects.all()
    }
    return render(request, 'entidades.html', context)

@login_required(login_url=reverse_lazy('login'))
def elimina_ofuscacion(request, pk):
    recurso = get_object_or_404(Ofuscacion, pk=pk)
    recurso.delete()
    return redirect('ofuscaciones')

class ActualizaOfuscacion(LoginRequiredMixin, UpdateView):
    model = Ofuscacion
    template_name = 'actualiza_ofuscacion.html'
    success_url = reverse_lazy('ofuscaciones')
    fields = ('nombre', 'regex')
    
class NuevaOfuscacion(LoginRequiredMixin, CreateView):
    model = Ofuscacion
    template_name = 'nueva_ofuscacion.html'
    success_url = reverse_lazy('ofuscaciones')
    fields = ('nombre', 'regex')
    
@login_required(login_url=reverse_lazy('login'))
def elimina_entidad(request, pk):
    entidad = get_object_or_404(Entidades, pk=pk)
    entidad.delete()
    return redirect('entidades')

class ActualizaEntidad(LoginRequiredMixin, UpdateView):
    model = Entidades
    template_name = 'actualiza_entidad.html'
    success_url = reverse_lazy('entidades')
    fields = ('nombre', 'formularios', 'clasificacion',)

    def get_form(self, form_class=None):
        form = super(ActualizaEntidad, self).get_form(form_class)
        form.fields['formularios'].required = False
        form.fields['clasificacion'].required = False
        form.fields['formularios'].widget = Textarea()
        form.fields['formularios'].label = 'Formularios de abuso (separados por un salto de línea)'
        return form
    
class NuevaEntidad(LoginRequiredMixin, CreateView):
    model = Entidades
    template_name = 'nueva_entidad.html'
    success_url = reverse_lazy('entidades')
    fields = ('nombre', 'formularios', 'clasificacion',)

    def get_form(self, form_class=None):
        form = super(NuevaEntidad, self).get_form(form_class)
        form.fields['formularios'].required = False
        form.fields['clasificacion'].required = False
        form.fields['formularios'].widget = Textarea()
        form.fields['formularios'].label = 'Formularios de abuso (separados por un salto de línea)'
        return form

@login_required(login_url=reverse_lazy('login'))
def elimina_clasificacion(request, pk):
    c = get_object_or_404(Clasificacion_entidad, pk=pk)
    c.delete()
    return redirect('entidades')

class NuevaClasificacionEntidad(LoginRequiredMixin, CreateView):
    model = Clasificacion_entidad
    template_name = 'nueva_clasificacion.html'
    success_url = reverse_lazy('entidades')
    fields = ('nombre',)

    def get_form(self, form_class=None):
        form = super(NuevaClasificacionEntidad, self).get_form(form_class)
        form.fields['nombre'].label = 'Clasificación'
        return form

class ActualizaClasificacionEntidad(LoginRequiredMixin, UpdateView):
    model = Clasificacion_entidad
    template_name = 'actualiza_clasificacion.html'
    success_url = reverse_lazy('entidades')
    fields = ('nombre',)

    def get_form(self, form_class=None):
        form = super(ActualizaClasificacionEntidad, self).get_form(form_class)
        form.fields['nombre'].label = 'Clasificación'
        return form

class HomeView(View):
    def get(self,request, *args, **kwargs):
        return render(request, 'dashboard.html', {})

def obtener_dias():
    """
    Obtinene una lista de los últimos 7 días
    """
    dias = ['Lunes', 'Martes', 'Miércoles', 'Jueves', 'Viernes', 'Sábado', 'Domingo']
    hoy = timezone.localtime(timezone.now()).weekday() + 1
    return dias[hoy:] + dias[:hoy]

def delta_horas(td):
    return td.total_seconds() / 3600

class ChartData(APIView):
    
    def get(self, request, format=None):
        rand_color = randomcolor.RandomColor()
        urls = Url.objects.exclude(Q(deteccion='I')|Q(deteccion='N'))

        paisesU = urls.filter(reportado=False).exclude(dominio__pais=None).values(
            'dominio__pais').annotate(
                cuenta_pais=Count('dominio__pais'))
        paisesMU = MensajeURL.objects.exclude(pais=None).values('pais').annotate(
            cuenta_pais=Count('pais'))
        l = []
        for s in paisesMU:
            try:
                n = next(t['cuenta_pais'] for t in paisesU if
                         t['dominio__pais'] == s['pais'])
            except:
                n = 0
            l.append((s['pais'], s['cuenta_pais'] + n))
        for s in paisesU:
            try:
                n = next(1 for t in l if t[0] == s['dominio__pais'])
            except:                        
                l.append((s['dominio__pais'], s['cuenta_pais']))
        l.sort(key=lambda x:x[1], reverse=True)
        l = l[:5]
        top_paises_data = {
            "labels": [a[0] for a in l],
            "default": [a[1] for a in l]
        }

        hostingU = urls.filter(reportado=False).exclude(dominio__asn=None).values(
            'dominio__asn').annotate(cuenta_asn=Count('dominio__asn'))
        hostingMU = MensajeURL.objects.exclude(asn=None).values(
            'asn').annotate(cuenta_asn=Count('asn'))
        l = []
        for s in hostingMU:
            try:
                n = next(t['cuenta_asn'] for t in hostingU if
                         t['dominio__asn'] == s['asn'])
            except:
                n = 0
            l.append((s['asn'], s['cuenta_asn'] + n))
        for s in hostingU:
            try:
                n = next(1 for t in l if t[0] == s['dominio__asn'])
            except:
                l.append((s['dominio__asn'], s['cuenta_asn']))
        l.sort(key=lambda x:x[1], reverse=True)
        l = l[:5]
        top_hosting_data = {
            "labels": [a[0] for a in l],
            "default": [a[1] for a in l]
        }
        
        sitios_activos = urls.filter(timestamp_desactivado=None).count()
        sitios_reportados = MensajeURL.objects.all().count()
        sitios_detectados = urls.filter(reportado=False).count() + MensajeURL.objects.all().count()
        sitios_data = {
            'labels': ['Activos', 'Reportados', 'Detectados'],
            'default': [sitios_activos, sitios_reportados, sitios_detectados]
        }
        
        hoy_tiempo = timezone.localtime(timezone.now())
        top_sitios = urls.filter(timestamp_desactivado=None).annotate(
                                            tiempo_vida=(hoy_tiempo -
                                                F('timestamp_reactivacion'))).order_by(
                                                    '-tiempo_vida')[:5]
        top_sitios_data = {
            'labels': [x.url for x in top_sitios],
            'default': [delta_horas(x.tiempo_vida) for x in top_sitios]
        }
        
        sectoresU = urls.filter(~Q(entidades_afectadas=None),
                                ~Q(entidades_afectadas__clasificacion=None),
                                reportado=False).values(
                                    'entidades_afectadas__clasificacion__nombre').annotate(
                                        cuenta_sectores=Count(
                                            'entidades_afectadas__clasificacion__nombre'))
        sectoresMU = MensajeURL.objects.filter(~Q(entidades_afectadas=None),
                                               ~Q(entidades_afectadas__clasificacion=None)).values(
                                                   'entidades_afectadas__clasificacion__nombre').annotate(
                                                       cuenta_sectores=Count(
                                                           'entidades_afectadas__clasificacion__nombre'))
        x, y = [], []
        for s in sectoresMU:
            try:
                n = next(t['cuenta_sectores'] for t in sectoresU if
                         t['entidades_afectadas__clasificacion__nombre'] == s['entidades_afectadas__clasificacion__nombre'])
            except:
                n = 0
                x.append(s['entidades_afectadas__clasificacion__nombre'])
                y.append(s['cuenta_sectores'] + n)
        for s in sectoresU:
            if not s['entidades_afectadas__clasificacion__nombre'] in x:
                x.append(s['entidades_afectadas__clasificacion__nombre'])
                y.append(s['cuenta_sectores'])
        sectores_data = {
            "labels":  x,
            "default": y,
            "colores": rand_color.generate(count=len(x))
        }
        
        dias = obtener_dias()
        num_detecciones = [] 
        hoy = hoy_tiempo.date()
        for x in range(6, -1, -1):
            num_detecciones.append(urls.filter(reportado=False,
                timestamp_deteccion__date=hoy - datetime.timedelta(days=x),
            ).count() + MensajeURL.objects.filter(
                timestamp_deteccion__date=hoy - datetime.timedelta(days=x)).count())
        detecciones_data = {
            'labels': dias,
            'default': num_detecciones
        }

        entidadesU = urls.filter(reportado=False).exclude(entidades_afectadas=None).values(
            'entidades_afectadas__nombre').annotate(
                cuenta_entidades=Count('entidades_afectadas__nombre'))
        entidadesMU = MensajeURL.objects.exclude(entidades_afectadas=None).values(
            'entidades_afectadas__nombre').annotate(
                cuenta_entidades=Count('entidades_afectadas__nombre'))
        x, y = [], []
        for s in entidadesMU:
            try:
                n = next(t['cuenta_entidades'] for t in entidadesU if
                         t['entidades_afectadas__nombre'] == s['entidades_afectadas__nombre'])
            except:
                n = 0
                x.append(s['entidades_afectadas__nombre'])
                y.append(s['cuenta_entidades'] + n)
        for s in entidadesU:
            if not s['entidades_afectadas__nombre'] in x:
                x.append(s['entidades_afectadas__nombre'])
                y.append(s['cuenta_entidades'])
        entidades_data = {
            "labels":  x,
            "default": y,
            "colores": rand_color.generate(count=len(x))
        }

        tiempo_promedio_reporte = []
        tiempo_promedio_postreporte = []
        for x in range(6, -1, -1):
            tickets = MensajeURL.objects.filter(mensaje__timestamp__date=hoy - datetime.timedelta(days=x))
            tiempo_promedio_reporte.append(tickets.annotate(
                tiempo_reportado=F('mensaje__timestamp') - F('timestamp_creacion_sitio')).aggregate(
                    Avg('tiempo_reportado')).get('tiempo_reportado__avg', 0))
            tiempo_promedio_postreporte.append(tickets.filter(
                ~Q(timestamp_desactivado=None)).annotate(
                    tiempo_reportado=F('timestamp_desactivado') - F('mensaje__timestamp')).aggregate(
                        Avg('tiempo_reportado')).get('tiempo_reportado__avg', 0))                
        tiempo_reporte_data = {
            'default1': [delta_horas(x) if x else 0 for x in tiempo_promedio_reporte],
            'default2': [delta_horas(x) if x else 0 for x in tiempo_promedio_postreporte]
        }
        
        graphs = [top_paises_data, top_hosting_data, sitios_data, top_sitios_data,
                  sectores_data, detecciones_data, entidades_data, tiempo_reporte_data]
        return Response(graphs)

class DocumentView(LoginRequiredMixin, View):
    def get(self,request, *args, **kwargs):
        if request.method == 'POST':
            return render(request,'dashboard.html',{})
        else:
            return render(request,'generar_rep.html', {'form': GraficasForm()})

def agrega_imagen(fig, documento):
    try:
        a = ''.join(random.choice(string.ascii_uppercase + string.digits) for _ in range(8))
        path = '/tmp/%s.png' % a
        fig.savefig(path)
        plt.clf()
        documento.add_picture(path)
        os.remove(path)
    except Exception as e:
        log_x('Error: %s' % str(e), 'reportes.log')

def url_info(u, q, d):
    if u.captura and hasattr(u.captura, 'file'):
        d.add_picture(u.captura.file, width=Inches(4.0))
    q = d.add_paragraph("")
    q.add_run("URL: %s\n" % u.url)
    q.add_run("Identificador: %s\n" % u.identificador)
    q.add_run("IP: %s\n" % u.dominio.ip)
    q.add_run("Código: %s\n" % u.codigo_estado)
    q.add_run("Fecha de creación: %s\n" % u.timestamp_creacion)
    if u.activo_redirecciones:
        q.add_run("Fecha de activación: %s\n" % u.timestamp_reactivacion)
    q.add_run("Detección: %s\n" % u.get_deteccion_display)                    
    q.add_run("Título: %s\n" % u.titulo)
    q.add_run("Estado: %s\n" % u.estado)                  
    q.add_run("Entidades: %s\n" % u.entidades)                    
    q.add_run("Ofuscacion: %s\n" % u.ofuscaciones)
    q.add_run("Correos: %s\n" % u.dominio.correos_abuso)
    q.add_run("ISP: %s\n" % u.dominio.isp)
    q.add_run("País: %s\n\n" % u.dominio.pais.name)
    q.add_run("ASN: %s\n" % u.dominio.asn)
    q.add_run("Servidor: %s\n" % u.dominio.servidor)
    q.add_run("RIR: %s\n" % u.dominio.rir)
    q.add_run("Servidores DNS: %s\n" % u.dominio.servidores_dns)
    if u.timestamp_deteccion:
        q.add_run("Fecha de detección: %s\n" % u.timestamp_deteccion)
    if u.timestamp_desactivado:
        q.add_run("Fecha de desactivación: %s\n" % u.timestamp_desactivado)
    if u.hash_archivo:
        q.add_run("Hash MD5 de archivo: %s\n" % u.hash_archivo)
    if u.codigo >= 300 and u.codigo < 400:
        q.add_run("Redirección: %s\n" % u.redireccion)
        r = u.get_redireccion
        if r:
            q.add_run("Redirección final: %s\n" % r.url)
                     
@login_required(login_url=reverse_lazy('login'))
def createDoc(request):
    if request.method == 'POST':
        form = GraficasForm(request.POST)
        if form.is_valid():
            archivo = form.cleaned_data['archivo']
            sitios = form.cleaned_data['sitios']
            top_sitios = form.cleaned_data['top_sitios']
            sectores = form.cleaned_data['sectores']
            entidades = form.cleaned_data['entidades']
            detecciones = form.cleaned_data['detecciones']
            tiempo_reporte = form.cleaned_data['tiempo_reporte']
            top_paises = form.cleaned_data['top_paises']
            top_hosting = form.cleaned_data['top_hosting']
            urls_info = form.cleaned_data['urls']
            graficas = []
            rand_color = randomcolor.RandomColor()
            document = Document()
            document.add_heading('Reporte', 0)
            p = document.add_paragraph('Reporte elaborado por la herramienta ')
            p.add_run('SAAPM').bold = True
            inicio = form.cleaned_data['inicio']
            fin = form.cleaned_data['fin']
            urlsU = Url.objects.exclude(timestamp_deteccion=None).filter(
                reportado=False,
                timestamp_deteccion__date__gte=inicio,
                timestamp_deteccion__date__lte=fin
            )
            urlsMU = MensajeURL.objects.exclude(timestamp_deteccion=None).filter(
                timestamp_deteccion__date__gte=inicio,
                timestamp_deteccion__date__lte=fin
            )
            document.add_heading('Periodo',level=1)
            q = document.add_paragraph('De: ')
            q.add_run(str(inicio)).bold = True
            q.add_run('      ')
            q.add_run('A :  ')
            q.add_run(str(fin)).bold = True

            if sitios:
                sitios_activos = urlsU.filter(Q(timestamp_desactivado=None) |
                                              Q(timestamp_desactivado__date__lt=fin)).count() + \
                                              urlsMU.filter(Q(timestamp_desactivado=None) |
                                                            Q(timestamp_desactivado__date__lt=fin)).count()
                sitios_reportados = urlsMU.count()
                sitios_detectados = urlsU.count() + urlsMU.count()
                x = ['Activos', 'Reportados', 'Detectados']
                y = [sitios_activos, sitios_reportados, sitios_detectados]
                y_pos = np.arange(len(x))
                fig, ax = plt.subplots()
                ax.set_ylabel('Número de sitios')
                ax.bar(y_pos, y, align='center', alpha=0.5)
                plt.xticks(y_pos, x)
                ax.set_title('Estados de sitios phishing')
                agrega_imagen(fig, document)

            if top_sitios:
                hoy = timezone.localtime(timezone.now())
                finf = hoy if hoy.date() == fin else fin + timedelta(days=1)
                top_sitiosU = urlsU.filter(Q(timestamp_desactivado=None) |
                                           Q(timestamp_desactivado__date__lte=fin)).annotate(
                                               tiempo_vida=(
                                                   finf - F('timestamp_reactivacion')))
                tsU = [(x.url, delta_horas(x.tiempo_vida)) for x in top_sitiosU]
                top_sitiosMU = urlsMU.filter(Q(timestamp_desactivado=None) |
                                             Q(timestamp_desactivado__date__lte=fin)).annotate(
                                                 tiempo_vida=(
                                                   finf - F('timestamp_creacion_sitio')))
                tsMU = [(x.url.url, delta_horas(x.tiempo_vida)) for x in top_sitiosMU]
                top_sitios = (tsMU + tsU)
                top_sitios.sort(key=lambda x:x[1], reverse=True)
                print(top_sitios)
                top_sitios = top_sitios[:5]
                y = [x[0] for x in top_sitios]
                x = [x[1] for x in top_sitios]
                y_pos = np.arange(len(x))
                fig, ax = plt.subplots()
                fig.subplots_adjust(left=0.5)
                ax.set_xlabel('T (Horas)')
                ax.barh(y_pos, x, align='center', alpha=0.5)
                plt.yticks(y_pos, y)
                ax.set_title('Top 5 – Sitios phishing vs Tiempo de vida')
                agrega_imagen(fig, document)
                
            if sectores:
                sectoresU = urlsU.filter(~Q(entidades_afectadas=None),
                                         ~Q(entidades_afectadas__clasificacion=None)).values(
                                             'entidades_afectadas__clasificacion__nombre').annotate(
                                                 cuenta_sectores=Count(
                                                     'entidades_afectadas__clasificacion__nombre'))
                sectoresMU = urlsMU.filter(~Q(entidades_afectadas=None),
                                           ~Q(entidades_afectadas__clasificacion=None)).values(
                                               'entidades_afectadas__clasificacion__nombre').annotate(
                                                   cuenta_sectores=Count(
                                                       'entidades_afectadas__clasificacion__nombre'))
                x, y = [], []
                for s in sectoresMU:
                    try:
                        n = next(t['cuenta_sectores'] for t in sectoresU if
                                 t['entidades_afectadas__clasificacion__nombre'] == s['entidades_afectadas__clasificacion__nombre'])
                    except:
                        n = 0
                    x.append(s['entidades_afectadas__clasificacion__nombre'])
                    y.append(s['cuenta_sectores'] + n)
                for s in sectoresU:
                    if not s['entidades_afectadas__clasificacion__nombre'] in x:
                        x.append(s['entidades_afectadas__clasificacion__nombre'])
                        y.append(s['cuenta_sectores'])
                colores = rand_color.generate(count=len(x))
                fig, ax = plt.subplots()
                ax.pie(y, labels=x, colors=colores, autopct='%1.1f%%', startangle=90)
                ax.set_title('Sectores afectados')
                ax.axis('equal')
                agrega_imagen(fig, document)
                
            if entidades:
                entidadesU = urlsU.exclude(entidades_afectadas=None).values(
                    'entidades_afectadas__nombre').annotate(
                        cuenta_entidades=Count('entidades_afectadas__nombre'))
                entidadesMU = urlsMU.exclude(entidades_afectadas=None).values(
                    'entidades_afectadas__nombre').annotate(
                        cuenta_entidades=Count('entidades_afectadas__nombre'))
                x, y = [], []
                for s in entidadesMU:
                    try:
                        n = next(t['cuenta_entidades'] for t in entidadesU if
                                 t['entidades_afectadas__nombre'] == s['entidades_afectadas__nombre'])
                    except:
                        n = 0
                    x.append(s['entidades_afectadas__nombre'])
                    y.append(s['cuenta_entidades'] + n)
                for s in entidadesU:
                    if not s['entidades_afectadas__nombre'] in x:
                        x.append(s['entidades_afectadas__nombre'])
                        y.append(s['cuenta_entidades'])
                colores = rand_color.generate(count=len(x))
                fig, ax = plt.subplots()
                ax.pie(y, labels=x, colors=colores, autopct='%1.1f%%', startangle=90)
                ax.set_title('Entidades afectadas')
                ax.axis('equal')
                agrega_imagen(fig, document)

            if detecciones:
                ndias = (fin - inicio).days
                fechas = [inicio + datetime.timedelta(days=i) for i in range(ndias + 1)]
                y = []
                for d in fechas:
                    y.append(urlsU.filter(timestamp_deteccion__date=d).count() +
                             urlsMU.filter(timestamp_deteccion__date=d).count()
                    )
                x = [str(f) for f in fechas]
                y_pos = np.arange(len(x))
                fig, ax = plt.subplots()
                fig.subplots_adjust(bottom=0.2)
                ax.set_ylabel('Número de detecciones')
                ax.bar(y_pos, y, align='center', alpha=0.5)
                plt.xticks(y_pos, x, rotation=45)
                ax.set_title('Número de detecciones por fecha')
                agrega_imagen(fig, document)

            if tiempo_reporte:
                ndias = (fin - inicio).days
                fechas = [inicio + datetime.timedelta(days=i) for i in range(ndias + 1)]
                y = []
                x = [str(f) for f in fechas]
                tiempo_promedio_reporte = []
                tiempo_promedio_postreporte = []
                for d in fechas:
                    tickets = urlsMU.filter(mensaje__timestamp__date=d)
                    tiempo_promedio_reporte.append(tickets.annotate(
                        tiempo_reportado=F('mensaje__timestamp') - F('timestamp_creacion_sitio')).aggregate(
                            Avg('tiempo_reportado')).get('tiempo_reportado__avg', 0))
                    tiempo_promedio_postreporte.append(tickets.filter(
                        ~Q(timestamp_desactivado=None)).annotate(
                            tiempo_reportado=F('timestamp_desactivado') - F('mensaje__timestamp')).aggregate(
                                Avg('tiempo_reportado')).get('tiempo_reportado__avg', 0))
                y1 = [delta_horas(x) if x else 0 for x in tiempo_promedio_reporte]
                y2 = [delta_horas(x) if x else 0 for x in tiempo_promedio_postreporte]
                fig, ax = plt.subplots()
                fig.subplots_adjust(bottom=0.3)
                line1, = ax.plot(x, y1, linewidth=2,
                                 label='Tiempo promedio de reporte')
                line2, = ax.plot(x, y2, linewidth=2,
                                 label='Tiempo promedio de vida postreporte')
                plt.xticks(rotation=45)
                ax.set_ylabel('T (Horas)')
                ax.set_title('Tiempo promedio de reporte vs tiempo promedio de vida postreporte')
                ax.legend(loc='best')
                agrega_imagen(fig, document)

            if top_paises:                
                paisesU = urlsU.exclude(dominio__pais=None).values('dominio__pais').annotate(
                    cuenta_pais=Count('dominio__pais'))
                paisesMU = urlsMU.exclude(pais=None).values('pais').annotate(
                    cuenta_pais=Count('pais'))
                l = []
                for s in paisesMU:
                    try:
                        n = next(t['cuenta_pais'] for t in paisesU if
                                 t['dominio__pais'] == s['pais'])
                    except:
                        n = 0
                    l.append((s['pais'], s['cuenta_pais'] + n))
                for s in paisesU:
                    try:
                        n = next(1 for t in l if t[0] == s['dominio__pais'])
                    except:                        
                        l.append((s['dominio__pais'], s['cuenta_pais']))
                l.sort(key=lambda x:x[1], reverse=True)
                l = l[:10]
                x = [a[0] for a in l]
                y = [a[1] for a in l]
                fig, ax = plt.subplots()
                ax.set_ylabel('Número de sitios')
                y_pos = np.arange(len(x))
                ax.bar(y_pos, y, align='center', alpha=0.5)
                plt.xticks(y_pos, x)
                ax.set_title('Top 10 países que hospedan phishing')
                agrega_imagen(fig, document)
                
            if top_hosting:
                hostingU = urlsU.exclude(dominio__asn=None).values('dominio__asn').annotate(
                    cuenta_asn=Count('dominio__asn'))
                hostingMU = urlsMU.exclude(asn=None).values('asn').annotate(
                    cuenta_asn=Count('asn'))    
                l = []
                for s in hostingMU:
                    try:
                        n = next(t['cuenta__asn'] for t in hostingU if
                                 t['dominio__asn'] == s['asn'])
                    except:
                        n = 0
                    l.append((s['asn'], s['cuenta_asn'] + n))
                for s in hostingU:
                    try:
                        n = next(1 for t in l if t[0] == s['dominio__asn'])
                    except:                        
                        l.append((s['dominio__asn'], s['cuenta_asn']))
                l.sort(key=lambda x:x[1], reverse=True)
                l = l[:10]
                x = [a[0] for a in l]
                y = [a[1] for a in l]                
                fig, ax = plt.subplots()
                fig.subplots_adjust(bottom=0.5)
                ax.set_ylabel('Número de sitios')
                y_pos = np.arange(len(x))
                ax.bar(y_pos, y, align='center', alpha=0.5)
                plt.xticks(y_pos, x, rotation=70)
                ax.set_title('Top 10 servicios de hosting que hospedan phishing')
                agrega_imagen(fig, document)
                
            if urls_info:
                activas = urls_activas(urlsU)
                inactivas = urls_inactivas(urlsU)
                redirecciones = urls_redirecciones(urlsU)
                dominios = urls_dominios(urlsU)
                q.add_run("\n\nURLS NO REPORTADAS\n").bold = True
                q = document.add_paragraph("URLs analizadas: %d\n" % cuenta_urls(urlsU))
                q.add_run("URLs activas: %d\n" % len(activas))
                q.add_run("URLs inactivas: %d\n" % len(inactivas))
                q.add_run("URLs redirecciones: %d\n" % len(redirecciones))
                q.add_run("Dominios afectados: %d" % len(dominios))
                q = document.add_paragraph("")
                q.add_run("Entidades:\n").bold = True
                for e in urls_entidades(urlsU):
                    q.add_run("%s\n" % e)
                q = document.add_paragraph("")
                q.add_run("Dominios:\n").bold = True
                for e in dominios:
                    q.add_run("%s\n" % e)
                q = document.add_paragraph("")
                q.add_run("Países:\n").bold = True
                for e in urls_paises(urlsU):
                    q.add_run("%s\n" % e)

                q = document.add_paragraph("")         
                q.add_run("SITIOS ACTIVOS:\n").bold = True
                for u in activas:
                    url_info(u, q, document)
                q = document.add_paragraph("")
                q.add_run("SITIOS INACTIVOS:\n\n").bold = True
                for u in inactivas:
                    url_info(u, q, document)
                q = document.add_paragraph("")
                q.add_run("REDIRECCIONES:\n\n").bold = True
                for u in redirecciones:
                    url_info(u, q, document)
                q = document.add_paragraph("")
                q.add_run("URLS REPORTADAS:\n\n").bold = True
                for u in urlsMU:
                    q = document.add_paragraph("")
                    q.add_run("URL: %s\n" % u.url)
                    q.add_run("Fecha de activación: %s\n" % u.timestamp_creacion_sitio)
                    if u.timestamp_desactivado:
                        q.add_run("Fecha de desactivación: %s\n" % u.timestamp_desactivado)
                    q.add_run("Fecha de detección: %s\n" % u.timestamp_deteccion)
                    q.add_run("Entidades afectadas: %s\n" % u.entidades)
                    q.add_run("País: %s\n" % u.pais)
                    q.add_run("ASN: %s\n" % u.asn)
                    
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = 'attachment; filename=%s.docx' % archivo
            document.save(response)
            return response
        else:
            return render(request,'generar_rep.html', {'form': GraficasForm()})

def entrada(request):
    if request.method == 'POST':
        if request.POST.get("boton_correo"):
            form = CorreoForm(request.POST)
            if form.is_valid():
                c = form.cleaned_data['correo']
                resultados, urls, headers, archivos, error = parsecorreo(c)
                sitios = []
                if len(urls) > 0:
                    sitios = verifica_urls(urls, None, False)
                context = context_reporte(sitios)                    
                context['resultados'] = resultados
                context['urls'] = urls
                context['headers'] = headers
                context['archivos'] = archivos
                context['error'] = error
                return render(request, 'entrada_resultados.html', context)
        elif request.POST.get("boton_archivo") and request.FILES.get('file', None):
            form = CorreoArchivoForm(request.POST)
            f = request.FILES['file'].read().decode('utf-8', 'ignore')
            name = request.FILES['file'].name
            urls = []
            resultados, urls, headers, archivos, error = parsecorreo(f)
            context = {}
            if len(urls) > 0:
                sitios = verifica_urls(urls, None, False)
                context = context_reporte(sitios)
            context['resultados'] = resultados
            context['urls'] = urls
            context['headers'] = headers
            context['archivos'] = archivos
            context['error'] = error
            return render(request, 'entrada_resultados.html', context)
    form1 = CorreoForm()
    form2 = CorreoArchivoForm()
    return render(request, 'entrada.html', {'form1': form1, 'form2': form2})

class TicketView(LoginRequiredMixin, DetailView):

    model = Mensaje
    template_name = 'ticket.html'
